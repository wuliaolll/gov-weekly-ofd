"""
文档生成模块 — DOCX + OFD
按照公文格式模板生成文件
"""
from __future__ import annotations

import io
import os
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

# ============ 字体注册 ============

import platform as _platform

# 字体搜索路径（跨平台）
if _platform.system() == "Windows":
    _FONT_DIRS = [
        Path("C:/Windows/Fonts"),
        Path.home() / "AppData/Local/Microsoft/Windows/Fonts",
    ]
else:
    # Linux / macOS — 常见系统字体目录 + 项目自带 fonts/ 目录
    _FONT_DIRS = [
        Path(__file__).parent / "fonts",
        Path.home() / ".fonts",
        Path("/usr/share/fonts"),
        Path("/usr/share/fonts/truetype"),
        Path("/usr/local/share/fonts"),
    ]


def _find_font(candidates: list[str]) -> str | None:
    for name in candidates:
        for d in _FONT_DIRS:
            p = d / name
            if p.exists():
                return str(p)
    return None


# 方正小标宋简体 — 标题
FZXBS_PATH = _find_font(["FZXBSJW.TTF", "fzxbsjw.ttf", "FZXBSFW.TTF", "方正小标宋简体.ttf", "方正小标宋_GBK.TTF"])

# 仿宋_GB2312 — 正文（优先 GB2312 版本）
FS_PATH = _find_font(["标准仿宋体简.TTF", "仿宋_GB2312.ttf", "FSGB2312.TTF", "simfang.ttf", "SIMFANG.TTF", "fangsong.ttf"])

# Times New Roman
TNR_PATH = _find_font(["times.ttf", "TIMES.TTF", "Times.ttf"])

# 注册 ReportLab 字体
_fonts_registered = False


def _register_reportlab_fonts():
    global _fonts_registered
    if _fonts_registered:
        return
    if FZXBS_PATH:
        pdfmetrics.registerFont(TTFont("FZXiaoBiaoSong", FZXBS_PATH))
        from reportlab.lib.fonts import addMapping
        addMapping("FZXiaoBiaoSong", 0, 0, "FZXiaoBiaoSong")
    if FS_PATH:
        pdfmetrics.registerFont(TTFont("FangSong", FS_PATH))
        from reportlab.lib.fonts import addMapping
        addMapping("FangSong", 0, 0, "FangSong")
    if TNR_PATH:
        pdfmetrics.registerFont(TTFont("TimesNewRoman", TNR_PATH))
        from reportlab.lib.fonts import addMapping
        addMapping("TimesNewRoman", 0, 0, "TimesNewRoman")
    _fonts_registered = True


# ============ 标题智能折行 ============

# 版心宽度: 21cm - 3.2cm - 3.2cm = 14.6cm; 标题 22pt，每 CJK 字符约 22pt 宽
# 14.6cm ≈ 414pt → 414/22 ≈ 18.8; 实测 Word 中 18 个全角字符刚好一行
_TITLE_MAX_CHARS = 18.0


def _char_width(ch: str) -> float:
    """CJK 全角 = 1，ASCII 半角 ≈ 0.5"""
    return 0.5 if ord(ch) < 128 else 1.0


def _text_width(s: str) -> float:
    return sum(_char_width(c) for c in s)


def _clean_title_text(title: str) -> str:
    """清理网页标题中的隐藏字符，保留可见空格与换行。"""
    if not title:
        return ""

    # 1. 去除所有换行变体（HTML 结构噪声，不是有含义的标题换行）。
    normalized = (
        title.replace("\r\n", "")
        .replace("\r", "")
        .replace("\n", "")
        .replace("\u2028", "")   # line separator
        .replace("\u2029", "")   # paragraph separator
        .replace("\u0085", "")   # NEL
    )
    # 2. 将所有 Unicode 空格类（Zs）统一为普通空格。
    #    覆盖范围：U+00A0 NBSP、U+2002 en-space、U+2003 em-space、
    #    U+202F narrow-NBSP、U+205F math-space、U+3000 全角空格等。
    normalized = "".join(
        " " if unicodedata.category(ch) == "Zs" else ch
        for ch in normalized
    )

    # 去除零宽/方向控制等隐藏字符，不影响普通空格和换行。
    hidden_chars = {
        "\u00AD",  # soft hyphen
        "\u034F",  # combining grapheme joiner
        "\u061C",  # arabic letter mark
        "\u180E",  # mongolian vowel separator
        "\u200B", "\u200C", "\u200D",  # zero-width chars
        "\u200E", "\u200F",  # lrm/rlm
        "\u2060",  # word joiner
        "\uFEFF",  # zero-width no-break space / bom
        # 网页复制噪声里常见的“圆圈样”占位符，需直接剔除。
        "\u00B0",  # degree sign: °
        "\u02DA",  # ring above: ˚
        "\u2218",  # ring operator: ∘
        "\u25E6",  # white bullet: ◦
        "\u25CB",  # white circle: ○
        "\u25CC",  # dotted circle: ◌
        # 回车/换行的可视化符号（不是实际换行），网页复制时可能混入。
        "\u21B5",  # carriage return arrow: ↵
        "\u23CE",  # return symbol: ⏎
        "\u240A",  # symbol for line feed: ␊
        "\u240D",  # symbol for carriage return: ␍
    }
    hidden_chars.update(chr(c) for c in range(0x202A, 0x202F))  # bidi embedding controls
    hidden_chars.update(chr(c) for c in range(0x2066, 0x206A))  # bidi isolate controls

    cleaned = "".join(ch for ch in normalized if ch not in hidden_chars)
    # 兜底剔除控制符（保留普通空格和换行）。
    cleaned = "".join(
        ch for ch in cleaned
        if ch == " " or unicodedata.category(ch) not in {"Cc", "Cf"}
    )

    # 5. 合并连续空格，去除首尾空格。
    return re.sub(r" {2,}", " ", cleaned).strip()


def _smart_title_lines(title: str, max_width: float = _TITLE_MAX_CHARS) -> list[str]:
    """基于语义断点的标题智能折行。

    1. jieba 分词 + 词性标注，合并短语原子单元
    2. 识别从句边界（动宾短语完成处）给予奖励
    3. 用 DP 在词边界处寻找语义最优的分行方案
       - 评分 = 断点语义惩罚之和（越低越好）
       - 轻微均匀性偏好作为 tiebreaker
    """
    title = _clean_title_text(title)

    total = _text_width(title)
    if total <= max_width:
        return [title]

    # --- 空格分隔的标题：直接按空格拆分短语，再合并成最少行 ---
    if ' ' in title:
        import re as _re
        segments = [s for s in _re.split(r'[ ]+', title) if s]
        # 所有片段都不超宽才走快速路径
        if all(_text_width(s) <= max_width for s in segments):
            import math as _m
            n_seg = len(segments)
            seg_widths = [_text_width(s) for s in segments]
            min_lines = _m.ceil(sum(seg_widths) / max_width)
            # DP: dp[i] = (min_lines, min_max_deviation, break_points) 前 i 个片段
            _INF = float('inf')
            dp = [(_INF, _INF, [])] * (n_seg + 1)
            dp[0] = (0, 0.0, [])
            for i in range(1, n_seg + 1):
                for j in range(i):
                    # segments[j..i-1] 合成一行
                    line_w = sum(seg_widths[j:i])
                    if line_w > max_width:
                        continue
                    prev_lines, prev_dev, prev_bps = dp[j]
                    new_lines = prev_lines + 1
                    new_dev = prev_dev + line_w  # 临时存总宽，最后算方差
                    if new_lines < dp[i][0] or (new_lines == dp[i][0]):
                        # 需要比较方差
                        pass
                    dp_candidate = (new_lines, prev_bps + [i])
                    # 简化：先找最少行方案，再在最少行方案中选最均匀的
                    pass
            # 简化实现：枚举所有分行方案（片段数通常 ≤5）
            best_lines = None
            best_score = (_INF, _INF)
            def _try_partition(seg_idx, current_lines):
                nonlocal best_lines, best_score
                if seg_idx == n_seg:
                    n = len(current_lines)
                    if n == 0:
                        return
                    widths = [_text_width(l) for l in current_lines]
                    mean_w = sum(widths) / n
                    variance = sum((w - mean_w) ** 2 for w in widths)
                    score = (n, variance)
                    if score < best_score:
                        best_score = score
                        best_lines = list(current_lines)
                    return
                for end in range(seg_idx + 1, n_seg + 1):
                    line = ' '.join(segments[seg_idx:end])
                    if _text_width(line) > max_width:
                        break
                    current_lines.append(line)
                    _try_partition(end, current_lines)
                    current_lines.pop()
            _try_partition(0, [])
            if best_lines:
                return best_lines
        # 片段超宽或无空格标题 → 落入 jieba 语义 DP

    # jieba DP 语义折行（有无空格都走此路径）
    space_positions: set[int] = set()

    import math
    import jieba
    import jieba.posseg as pseg

    # 注册政务文本常见复合词，避免被拆分
    for cw in ('科技创新', '产业创新', '城市更新', '融合发展',
               '高质量发展', '春季学期', '秋季学期',
               '营商环境', '经济社会', '投融资',
               '城建领域', '体制改革', '专题会议', '常务会议',
               '深入推进', '扎实推进', '统筹推进', '科学务实'):
        jieba.add_word(cw)

    # 词性标注
    word_pos_list = list(pseg.cut(title))
    words = [w for w, _ in word_pos_list]
    flags = [f for _, f in word_pos_list]

    # ---- Phase 1: 短语合并 ----
    _MOD_POS = {'a', 'ad', 'd', 'vd'}
    merged_words = []
    merged_flags = []
    i = 0
    while i < len(words):
        # 短修饰词(≤2) + 短动词(≤2) → 复合动词 (牢固+树立, 坚定+践行)
        if (i + 1 < len(words)
                and len(words[i]) <= 2 and len(words[i + 1]) <= 2
                and flags[i] in _MOD_POS
                and flags[i + 1].startswith('v')):
            merged_words.append(words[i] + words[i + 1])
            merged_flags.append(flags[i + 1])
            i += 2
        # 短内容词(≤2) + 短修饰词(≤2) → 并列修饰语 (科学+务实)
        # 排除：前一个词是动词（说明当前词是宾语，不应与后面合并）
        elif (i + 1 < len(words)
              and len(words[i]) <= 2 and len(words[i + 1]) <= 2
              and flags[i] in {'n', 'a', 'ad', 'vn'}
              and flags[i + 1] in _MOD_POS
              and not (i > 0 and flags[i - 1].startswith('v'))):
            merged_words.append(words[i] + words[i + 1])
            merged_flags.append(flags[i])
            i += 2
        else:
            merged_words.append(words[i])
            merged_flags.append(flags[i])
            i += 1

    words = merged_words
    flags = merged_flags

    # ---- Phase 2: 为每个词边界计算语义断点分数 ----
    _NOUN_POS = {'n', 'vn', 'ns', 'nt', 'nz', 'l'}
    _STICKY_POS = {'d', 'a', 'ad', 'p', 'c', 'f', 'r', 'm', 'q', 'vd'}

    break_positions = []  # [(char_pos, penalty)]
    pos = 0
    for i, w in enumerate(words):
        pos += len(w)
        # 空格 token 是自然短语边界，记录其结束位置
        if w == " ":
            space_positions.add(pos)
        if i < len(words) - 1:
            flag_i = flags[i]
            flag_next = flags[i + 1]
            penalty = 0.0

            # --- 原始空格位置是优质断点（标题中空格表示短语分隔） ---
            if pos in space_positions:
                penalty -= 15.0

            # --- 从句边界奖励 ---
            # 名词完成动宾短语（前面近处有动词），后面开始新动词短语
            # 单字名词如"时""上"太短，不足以标识从句边界
            if flag_i in _NOUN_POS and len(w) >= 2:
                has_verb_before = any(
                    flags[j].startswith('v')
                    for j in range(max(0, i - 4), i)
                )
                if has_verb_before:
                    starts_new = False
                    if flag_next.startswith('v'):
                        starts_new = True
                    elif flag_next in _MOD_POS:
                        for j in range(i + 2, min(i + 5, len(words))):
                            if flags[j].startswith('v'):
                                starts_new = True
                                break
                            if flags[j] not in _MOD_POS | {'n'}:
                                break
                    if starts_new:
                        # 前一个词是副词(d) → 在修饰链内（如"更加 科学务实 推动"），
                        # 不是真正的从句边界，不给奖励
                        if not (i > 0 and flags[i - 1] == 'd'):
                            penalty -= 15.0

            # --- 惩罚：不良断点 ---
            # 修饰/功能词做行尾
            if flag_i in _STICKY_POS:
                penalty += 12.0
            # 单字功能词做行尾
            if len(w) == 1 and flag_i in {'c', 'p', 'u', 'f', 'r'}:
                penalty += 8.0
            # 动词后跟它的修饰词/宾语，不宜拆开（如"踐行正确政绩观")
            if flag_i.startswith('v') and flag_i != 'vn':
                if flag_next in _NOUN_POS | {'ad', 'a'}:
                    penalty += 5.0
            # 连续短动词（复合动词）
            if (flag_i.startswith('v') and flag_next.startswith('v')
                    and len(w) <= 2 and len(words[i + 1]) <= 2):
                penalty += 5.0

            break_positions.append((pos, penalty))

    # ---- Phase 3: DP 搜索最优分行 ----
    n_lines = math.ceil(total / max_width)
    INF = float("inf")

    best_result = None
    best_score = INF

    # 只尝试最少行数和多一行（多一行可利用更多从句边界）
    for target_n in range(n_lines, min(n_lines + 2, len(words) + 1)):
        if target_n < 1:
            continue
        target_w = total / target_n
        extra_line_penalty = (target_n - n_lines) * 6.0

        def search(start_idx: int, lines_left: int, memo: dict,
                   _tw: float = target_w) -> tuple:
            key = (start_idx, lines_left)
            if key in memo:
                return memo[key]

            rw = _text_width(title[start_idx:])

            if lines_left == 1:
                if rw <= max_width:
                    # 轻微均匀性 + 过短行惩罚
                    score = 0.15 * (rw - _tw) ** 2 + max(0, 5 - rw) * 2.0
                    memo[key] = (score, [])
                else:
                    memo[key] = (INF, [])
                return memo[key]

            best_s = INF
            best_bp = []

            for bp, bp_penalty in break_positions:
                if bp <= start_idx:
                    continue
                line_w = _text_width(title[start_idx:bp])
                if line_w > max_width:
                    break
                if line_w < 3:
                    continue

                # 轻微均匀性 + 过短行惩罚
                line_score = 0.15 * (line_w - _tw) ** 2 + max(0, 5 - line_w) * 2.0

                sub_score, sub_bps = search(bp, lines_left - 1, memo)
                if sub_score >= INF:
                    continue
                score = bp_penalty + line_score + sub_score
                if score < best_s:
                    best_s = score
                    best_bp = [bp] + sub_bps

            memo[key] = (best_s, best_bp)
            return memo[key]

        score, bps = search(0, target_n, {})
        score += extra_line_penalty
        if score < best_score:
            best_score = score
            best_result = bps

    if not best_result:
        # 兜底：按词语贪心填充
        lines, cur = [], ""
        for w in words:
            if _text_width(cur + w) > max_width and cur:
                lines.append(cur)
                cur = w
            else:
                cur += w
        if cur:
            lines.append(cur)
        return lines

    lines = []
    prev = 0
    for bp in best_result:
        lines.append(title[prev:bp])
        prev = bp
    lines.append(title[prev:])
    return lines


# ============ DOCX 生成 ============

def generate_docx(title: str, content: str, output_path: str, paragraphs: list[dict] | None = None):
    """
    按公文格式生成 DOCX 文件

    页面: A4, 页边距上3.7cm 下3.0cm 左3.2cm 右3.2cm
    标题: 方正小标宋简体, 二号(22pt), 居中, 第二行, 与正文空一行
    正文: 仿宋GB2312, 小二号(18pt), 首行缩进2字符, 行距固定32磅, 两端对齐
    数字英文: Times New Roman
    颜色: 黑色
    无页码, 无图片
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)

    # 删除默认段落
    if doc.paragraphs:
        doc.paragraphs[0].clear()

    # 标题上方空一行
    top_spacer = doc.add_paragraph()
    top_spacer.paragraph_format.space_before = Pt(0)
    top_spacer.paragraph_format.space_after = Pt(0)
    _set_line_spacing_fixed(top_spacer, 32)

    # 标题（智能折行：每行独立段落，均居中，行间距固定 32pt）
    # 每行用独立 <w:p> 而非 <w:br/>，使 Word 中每行都显示为"正常段落换行"（←/¶）。
    title_lines = _smart_title_lines(title)
    for line in title_lines:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_spacing_fixed(title_para, 32)
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        # 拆分中文和英文/数字，分别设置字体
        for text, is_ascii in _split_cn_en(line):
            title_run = title_para.add_run(text)
            title_run.font.size = Pt(22)  # 二号
            if is_ascii:
                _set_font_name(title_run, "Times New Roman", "Times New Roman")
            else:
                _set_font_name(title_run, "方正小标宋简体", "FZXiaoBiaoSong")
            title_run.font.color.rgb = None  # 黑色

    # 标题与正文间空一行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    _set_line_spacing_fixed(spacer, 32)

    # 正文段落
    if paragraphs:
        para_list = paragraphs
    else:
        para_list = [{"text": p.strip(), "align": "left"} for p in content.split("\n") if p.strip()]
    prev_align_docx = None
    for para_info in para_list:
        para_text = para_info["text"] if isinstance(para_info, dict) else para_info
        para_align = para_info.get("align", "left") if isinstance(para_info, dict) else "left"

        # 居中段落后接正文时插入空行
        if prev_align_docx == "center" and para_align != "center":
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(0)
            _set_line_spacing_fixed(sp, 32)

        p = doc.add_paragraph()
        if para_align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_first_line_indent_chars(p, 2)  # 严格2字符缩进
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        _set_line_spacing_fixed(p, 32)

        # 拆分中文和英文/数字，分别设置字体
        segments = _split_cn_en(para_text)
        for text, is_ascii in segments:
            run = p.add_run(text)
            run.font.size = Pt(18)  # 小二号
            if is_ascii:
                _set_font_name(run, "Times New Roman", "Times New Roman")
            else:
                _set_font_name(run, "仿宋_GB2312", "FangSong_GB2312")
            run.font.bold = False
            run.font.color.rgb = None  # 黑色

        prev_align_docx = para_align

    doc.save(output_path)


def _set_line_spacing_fixed(paragraph, pt_value: int):
    """设置固定行距"""
    pf = paragraph.paragraph_format
    pf.line_spacing = Pt(pt_value)
    # 确保是固定值
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = pPr.makeelement(qn("w:spacing"), {})
        pPr.append(spacing)
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:line"), str(int(pt_value * 20)))  # 转为 twips


def _set_font_name(run, cn_name: str, en_name: str):
    """设置中英文字体名"""
    run.font.name = en_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), en_name)
    rFonts.set(qn("w:hAnsi"), en_name)
    rFonts.set(qn("w:eastAsia"), cn_name)
    rFonts.set(qn("w:cs"), en_name)


def _set_first_line_indent_chars(paragraph, chars: int):
    """用 OOXML w:firstLineChars 严格设置首行缩进字符数（Word「2字符」缩进标准写法）"""
    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = pPr.makeelement(qn("w:ind"), {})
        pPr.append(ind)
    # w:firstLineChars 单位: 1/100 字符，2字符 = 200
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    # w:firstLine 单位: twips（1/20 pt），18pt * 2 = 36pt = 720 twips（作为回退值）
    ind.set(qn("w:firstLine"), str(int(18 * 2 * 20)))


def _split_cn_en(text: str) -> list[tuple[str, bool]]:
    """将文本拆分为中文段和ASCII段"""
    segments = []
    current = ""
    current_is_ascii = False

    for ch in text:
        is_ascii = ord(ch) < 128 and ch not in " \t"
        # 空格/制表符跟随前一段
        if ch in " \t":
            current += ch
            continue
        if not current:
            current = ch
            current_is_ascii = is_ascii
        elif is_ascii == current_is_ascii:
            current += ch
        else:
            segments.append((current, current_is_ascii))
            current = ch
            current_is_ascii = is_ascii

    if current:
        segments.append((current, current_is_ascii))

    return segments


def _rl_mixed_font_text(text: str, cn_font: str, en_font: str | None) -> str:
    """为 ReportLab Paragraph 生成混合字体的 XML 标记文本。
    中文部分用 cn_font，ASCII 数字/字母部分用 en_font。
    在中西文边界插入一个 TNR 空格，模拟 Word「调整中西文间距」行为。
    文本需先经过 _xml_escape。"""
    if not en_font:
        return _xml_escape(text)
    segments = _split_cn_en(text)
    parts = []
    for i, (seg_text, is_ascii) in enumerate(segments):
        safe = _xml_escape(seg_text)
        if is_ascii:
            prev_is_cjk = i > 0 and not segments[i - 1][1]
            next_is_cjk = i < len(segments) - 1 and not segments[i + 1][1]
            # 前有 CJK → 插入一个 TNR 空格（约 1/4 em），模拟中西文间距
            prefix = f'<font name="{en_font}"> </font>' if prev_is_cjk else ""
            suffix = f'<font name="{en_font}"> </font>' if next_is_cjk else ""
            parts.append(f'{prefix}<font name="{en_font}">{safe}</font>{suffix}')
        else:
            parts.append(safe)
    return ''.join(parts)


# ============ OFD 生成 ============

def generate_ofd(title: str, content: str, output_path: str, paragraphs: list[dict] | None = None):
    """
    生成 OFD 文件：先用 ReportLab 生成 PDF，再用 easyofd 转换为 OFD
    """
    _register_reportlab_fonts()

    # 补充 ReportLab CJK 禁则字符表（缺少常用中文标点）
    import reportlab.platypus.paragraph as _rl_para
    import reportlab.lib.textsplit as _rl_ts
    _extra_no_start = '\uff0c\uff1b\uff1a\u201d\u2019\u300b\u3011\uff5d\u2026\u2014\u2013\uff5e'
    if '\uff0c' not in _rl_para.ALL_CANNOT_START:
        _rl_para.ALL_CANNOT_START += _extra_no_start
    if '\uff0c' not in _rl_ts.ALL_CANNOT_START:
        _rl_ts.ALL_CANNOT_START += _extra_no_start

    # 确定可用字体
    title_font = "FZXiaoBiaoSong" if FZXBS_PATH else "FangSong" if FS_PATH else "Helvetica"
    body_font = "FangSong" if FS_PATH else "Helvetica"

    # 构建 PDF 样式
    title_style = ParagraphStyle(
        "GovTitle",
        fontName=title_font,
        fontSize=22,
        leading=32,
        alignment=TA_CENTER,
        textColor="black",
        spaceAfter=0,
        spaceBefore=0,
        wordWrap='CJK',
    )

    body_style = ParagraphStyle(
        "GovBody",
        fontName=body_font,
        fontSize=18,
        leading=32,
        alignment=TA_JUSTIFY,
        firstLineIndent=18 * 2,  # 严格 2 字符：CJK 全角宽 = fontSize，2字符 = 36pt
        textColor="black",
        spaceAfter=0,
        spaceBefore=0,
        wordWrap='CJK',
    )

    # 生成 PDF 到内存
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        topMargin=3.7 * cm,
        bottomMargin=3.0 * cm,
        leftMargin=3.2 * cm,
        rightMargin=3.2 * cm,
    )

    story = []

    # 标题上方空一行
    story.append(Spacer(1, 32))

    # 标题（智能折行）
    en_font = "TimesNewRoman" if TNR_PATH else None
    title_lines = _smart_title_lines(title)
    safe_title = "<br/>".join(
        _rl_mixed_font_text(line, title_font, en_font) for line in title_lines
    )
    story.append(Paragraph(safe_title, title_style))

    # 标题与正文间空一行
    story.append(Spacer(1, 32))

    # 居中样式（用于居中段落）
    center_body_style = ParagraphStyle(
        "GovBodyCenter",
        fontName=body_font,
        fontSize=18,
        leading=32,
        alignment=TA_CENTER,
        firstLineIndent=0,
        textColor="black",
        spaceAfter=0,
        spaceBefore=0,
        wordWrap='CJK',
    )

    # 正文段落
    if paragraphs:
        para_list = paragraphs
    else:
        para_list = [{"text": p.strip(), "align": "left"} for p in content.split("\n") if p.strip()]
    prev_align = None
    for para_info in para_list:
        para_text = para_info["text"] if isinstance(para_info, dict) else para_info
        para_align = para_info.get("align", "left") if isinstance(para_info, dict) else "left"
        # 居中段落后接正文段落时，插入一个空行
        if prev_align == "center" and para_align != "center":
            story.append(Spacer(1, 32))
        safe_text = _rl_mixed_font_text(para_text, body_font, en_font)
        style = center_body_style if para_align == "center" else body_style
        story.append(Paragraph(safe_text, style))
        prev_align = para_align

    doc.build(story)

    # PDF → OFD
    pdf_bytes = pdf_buffer.getvalue()
    _pdf_to_ofd(pdf_bytes, output_path)


def _xml_escape(text: str) -> str:
    """转义文本中的 XML 特殊字符"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


def _pdf_to_ofd(pdf_bytes: bytes, output_path: str):
    """用 pdf2ofd skill 的 PDF2OFDConverter 将 PDF 转为 OFD（字符级精度）"""
    from pdf2ofd import PDF2OFDConverter
    converter = PDF2OFDConverter()
    ofd_bytes = converter.convert(pdf_bytes)
    with open(output_path, "wb") as f:
        f.write(ofd_bytes)


if __name__ == "__main__":
    # 快速测试
    test_title = "李殿勋主持召开省政府常务会议研究推进我省美丽中国先行区建设等工作"
    test_content = (
        "3月30日，省委副书记、省长李殿勋主持召开省政府常务会议，传达学习习近平总书记近期重要讲话精神；"
        "研究推进我省美丽中国先行区建设、自然保护地整合优化等工作；听取全省安全生产和森林防灭火工作情况汇报。\n\n"
        "会议指出，要深入学习贯彻习近平生态文明思想，牢固树立绿水青山就是金山银山的理念，"
        "加快推进美丽中国先行区建设，为全国生态文明建设贡献湖北力量。"
    )

    os.makedirs("test_output", exist_ok=True)
    print("生成 DOCX...")
    generate_docx(test_title, test_content, "test_output/test.docx")
    print("生成 OFD...")
    generate_ofd(test_title, test_content, "test_output/test.ofd")
    print("完成！文件已生成到 test_output/")
